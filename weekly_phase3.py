"""
weekly_phase3.py
Phase 3: เติมข้อมูลส่วนที่ต้อง input เพิ่มเติม
- ผลการดำเนินงาน (Table 0 summary + Table 1 detailed) จาก data/construction_plan.xlsx
- ภาคผนวก 4 บุคลากร CM 8 วัน จาก data/cm_personnel.xlsx
"""

import os, io, copy
from datetime import date, timedelta
from typing import Optional
import openpyxl
from openpyxl.styles import Font, Alignment
from copy import deepcopy
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

_THIS_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_PLAN  = os.path.join(_THIS_DIR, "data", "construction_plan.xlsx")
DATA_CM    = os.path.join(_THIS_DIR, "data", "cm_personnel.xlsx")
TEMPLATE_APPENDIX4_XLSX = os.path.join(_THIS_DIR, "templates_weekly", "appendix4_cm_personnel.xlsx")
TEMPLATE_APPENDIX4_1_DOCX = os.path.join(_THIS_DIR, "templates_weekly", "appendix4_1_activities.docx")

# คงที่: สัดส่วนงาน 5 หมวดใหญ่ (% ของ budget)
PROGRESS_SUMMARY_CATEGORIES = [
    ("1", "งานเขื่อนป้องกันตลิ่ง",                    33.35),
    ("2", "งานปรับปรุงภูมิทัศน์",                      10.91),
    ("3", "งานทาง",                                  54.43),
    ("4", "งานเครื่องจักรและอุปกรณ์และงานอื่นๆ",       1.07),
    ("5", "ค่าใช้จ่ายพิเศษ งานก่อสร้าง",               0.24),
]

# Map สำหรับเดือน → ชื่อ sheet ใน cm_personnel.xlsx
_MONTH_SHEET_MAP = {
    1:  "มค",  2: "กพ",  3: "มีค", 4: "เมย", 5: "พค", 6: "มิย",
    7:  "กค",  8: "สค",  9: "กย",  10: "ตค", 11: "พย", 12: "ธค",
}


# ════════════════════════════════════════
# Read: ผลการดำเนินงาน (ผลweekly sheet)
# ════════════════════════════════════════

def lookup_week_number(start_date: date, end_date: date) -> Optional[int]:
    """หา week_no จริงจาก sheet 'แผน - ผล ประจำสัปดาห์' โดยจับคู่ from_date + to_date
    คืน int week_no หรือ None ถ้าไม่พบ
    """
    if not os.path.exists(DATA_PLAN):
        return None
    try:
        wb = openpyxl.load_workbook(DATA_PLAN, data_only=True)
        if "แผน - ผล ประจำสัปดาห์" not in wb.sheetnames:
            return None
        ws = wb["แผน - ผล ประจำสัปดาห์"]
        # cols: D(4)=from_date, E(5)=to_date, F(6)=week_no
        from datetime import datetime as _dt
        for r in ws.iter_rows(min_row=5, values_only=True):
            if len(r) < 6:
                continue
            d_from = r[3]; d_to = r[4]; wk = r[5]
            if isinstance(d_from, _dt) and isinstance(d_to, _dt):
                if d_from.date() == start_date and d_to.date() == end_date:
                    return int(wk) if wk is not None else None
        # ถ้าไม่ exact match ลอง fallback: หา row ที่ start_date อยู่ในช่วง [from, to]
        for r in ws.iter_rows(min_row=5, values_only=True):
            if len(r) < 6:
                continue
            d_from = r[3]; d_to = r[4]; wk = r[5]
            if isinstance(d_from, _dt) and isinstance(d_to, _dt):
                if d_from.date() <= start_date <= d_to.date():
                    return int(wk) if wk is not None else None
    except Exception as e:
        print(f"⚠️ lookup_week_number failed: {e}")
    return None


def read_progress_detail() -> list:
    """อ่าน sheet 'ผลweekly' → list ของ row dict
    คืน: [{"no": "1.1", "name": "...", "share": 2.134, "prev_cum": 0.15, "this_cum": 0.15, "note": ""}, ...]
    """
    if not os.path.exists(DATA_PLAN):
        return []
    wb = openpyxl.load_workbook(DATA_PLAN, data_only=True)
    if "ผลweekly" not in wb.sheetnames:
        return []
    ws = wb["ผลweekly"]
    rows = []
    # data starts at row 5, columns: B=ที่, C=งานที่ดำเนินการ, D=สัดส่วน%, E=สะสมก่อน, F=สะสมนี้, G=หมายเหตุ
    for r in ws.iter_rows(min_row=5, values_only=True):
        if r is None or len(r) < 6:
            continue
        no = r[1]
        name = r[2]
        if no is None and name is None:
            continue
        rows.append({
            "no":       str(no).strip() if no is not None else "",
            "name":     str(name).strip() if name is not None else "",
            "share":    float(r[3]) if r[3] not in (None, "") else None,
            "prev_cum": float(r[4]) if r[4] not in (None, "") else None,
            "this_cum": float(r[5]) if r[5] not in (None, "") else None,
            "note":     str(r[6]).strip() if len(r) > 6 and r[6] is not None else "",
        })
    return rows


def compute_progress_summary(detail_rows: list) -> list:
    """รวม cumulative % ตามหมวดใหญ่ 1-5
    คืน: [{"no": "1", "name": "งานเขื่อน...", "share": 33.35, "prev_cum": 7.17, "this_cum": 7.80}, ...]
    """
    summary = []
    for cat_no, cat_name, share in PROGRESS_SUMMARY_CATEGORIES:
        prev_cum = 0.0
        this_cum = 0.0
        for row in detail_rows:
            no = row["no"]
            # ตรงกับหมวดนี้: เริ่มต้นด้วย "{cat_no}." แต่ไม่ใช่ header (ที่ no=cat_no เปล่าๆ)
            if no.startswith(f"{cat_no}.") and "." in no:
                # นับเฉพาะ top-level x.y (1.1, 1.2, ...) ไม่นับ sub (1)(2)
                parts = no.split(".")
                if len(parts) == 2 and parts[0] == cat_no and parts[1].isdigit():
                    if row["prev_cum"] is not None:
                        prev_cum += row["prev_cum"]
                    if row["this_cum"] is not None:
                        this_cum += row["this_cum"]
        summary.append({
            "no":       cat_no,
            "name":     cat_name,
            "share":    share,
            "prev_cum": prev_cum,
            "this_cum": this_cum,
        })
    return summary


# ════════════════════════════════════════
# Fill: Tables 0, 1 ใน 04_project_details.docx
# ════════════════════════════════════════

def _set_cell(cell, text, font_size=14, bold=False, center=True):
    """ล้าง cell แล้วใส่ข้อความใหม่ (เลียนแบบ _set_cell_text จาก phase1)"""
    p = cell.paragraphs[0]
    for old_p in cell.paragraphs[1:]:
        old_p._element.getparent().remove(old_p._element)
    for r in p.runs:
        r._element.getparent().remove(r._element)
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(str(text))
    run.font.name = "TH SarabunIT๙"
    run.font.size = Pt(font_size)
    run.bold = bold
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for attr in ("ascii", "hAnsi", "cs"):
        rFonts.set(qn(f"w:{attr}"), "TH SarabunIT๙")


def _fmt_pct(v: Optional[float], decimals: int = 2) -> str:
    """format ตัวเลข % โดย trim trailing zeros เกินไป"""
    if v is None:
        return ""
    if v == 0:
        return "0.00"
    return f"{v:.{decimals}f}"


def fill_progress_summary_table(table, summary_rows: list):
    """เติม Table 0 (สรุปผลการดำเนินงาน) — 8 rows × 6 cols
    rows[0-1]: header
    rows[2-6]: 5 หมวด
    rows[7]: รวม
    """
    n = min(len(summary_rows), len(table.rows) - 3)  # หัก header 2 + รวม 1
    sum_share = sum_prev = sum_this = 0.0
    for i in range(n):
        row = table.rows[2 + i]
        s = summary_rows[i]
        _set_cell(row.cells[0], s["no"])
        _set_cell(row.cells[1], s["name"], center=False)
        _set_cell(row.cells[2], _fmt_pct(s["share"]))
        _set_cell(row.cells[3], _fmt_pct(s["prev_cum"]) + "%")
        _set_cell(row.cells[4], _fmt_pct(s["this_cum"]) + "%")
        _set_cell(row.cells[5], "")
        sum_share += s["share"]
        sum_prev  += s["prev_cum"]
        sum_this  += s["this_cum"]
    # รวม row (last)
    last = table.rows[-1]
    _set_cell(last.cells[0], "")
    _set_cell(last.cells[1], "รวม", bold=True)
    _set_cell(last.cells[2], _fmt_pct(sum_share), bold=True)
    _set_cell(last.cells[3], _fmt_pct(sum_prev) + "%", bold=True)
    _set_cell(last.cells[4], _fmt_pct(sum_this) + "%", bold=True)
    _set_cell(last.cells[5], "")


def _is_category_row(no: str) -> bool:
    """row ที่ no เป็นเลขเดียว เช่น "1", "2" → category header"""
    return no.isdigit() and "." not in no


def fill_progress_detail_table(table, detail_rows: list):
    """เติม Table 1 (สรุปผลการดำเนินงานละเอียด) — 55 rows × 6 cols
    Row[0-1]: header (skip)
    Row[2]: category template (มีสีพื้น C0E6F5) — ใช้สำหรับ category rows (1, 2, 3, ...)
    Row[3+]: data template (ไม่มีสีพื้น) — ใช้สำหรับ sub-rows (1.1, 1.2, (1), (2), ...)
    Last row: รวม template
    """
    if len(table.rows) < 5:
        return
    category_tpl = deepcopy(table.rows[2]._tr)  # category row (with shading)
    data_tpl     = deepcopy(table.rows[3]._tr)  # data row (no shading)
    sum_tpl      = deepcopy(table.rows[-1]._tr)

    # ลบ rows 2..end (เก็บแค่ 2 header rows)
    while len(table.rows) > 2:
        table.rows[-1]._element.getparent().remove(table.rows[-1]._element)

    sum_share = sum_prev = sum_this = 0.0
    sum_row_data = None
    for d in detail_rows:
        if "รวม" in d["name"]:
            sum_row_data = d
            continue
        # เลือก template ตาม row type
        is_cat = _is_category_row(d["no"])
        new_tr = deepcopy(category_tpl if is_cat else data_tpl)
        table._tbl.append(new_tr)
        new_row = table.rows[-1]
        _set_cell(new_row.cells[0], d["no"], bold=is_cat)
        _set_cell(new_row.cells[1], d["name"], center=False, bold=is_cat)
        _set_cell(new_row.cells[2], _fmt_pct(d["share"], 3) if d["share"] is not None else "", bold=is_cat)
        _set_cell(new_row.cells[3], _fmt_pct(d["prev_cum"]) if d["prev_cum"] is not None else "", bold=is_cat)
        _set_cell(new_row.cells[4], _fmt_pct(d["this_cum"]) if d["this_cum"] is not None else "", bold=is_cat)
        _set_cell(new_row.cells[5], d.get("note", ""), bold=is_cat)
        # accumulate top-level x.y subtotals
        no = d["no"]
        if "." in no and "(" not in no:
            parts = no.split(".")
            if len(parts) == 2 and parts[0].isdigit() and parts[1].isdigit():
                if d["share"]:    sum_share += d["share"]
                if d["prev_cum"]: sum_prev  += d["prev_cum"]
                if d["this_cum"]: sum_this  += d["this_cum"]

    # สรุปสุดท้าย "รวม" row
    new_tr = deepcopy(sum_tpl)
    table._tbl.append(new_tr)
    new_row = table.rows[-1]
    _set_cell(new_row.cells[0], "", bold=True)
    _set_cell(new_row.cells[1], "รวม", bold=True)
    if sum_row_data:
        _set_cell(new_row.cells[2], _fmt_pct(sum_row_data["share"]), bold=True)
        _set_cell(new_row.cells[3], _fmt_pct(sum_row_data["prev_cum"]), bold=True)
        _set_cell(new_row.cells[4], _fmt_pct(sum_row_data["this_cum"]), bold=True)
    else:
        _set_cell(new_row.cells[2], _fmt_pct(sum_share), bold=True)
        _set_cell(new_row.cells[3], _fmt_pct(sum_prev), bold=True)
        _set_cell(new_row.cells[4], _fmt_pct(sum_this), bold=True)
    _set_cell(new_row.cells[5], "", bold=True)


# ════════════════════════════════════════
# CM Personnel — read & fill
# ════════════════════════════════════════

def _sheet_name_for_month(month: int, year_be: int) -> list:
    """คืน list ของชื่อ sheet ที่อาจตรงกัน เช่น เดือน 4, ปี 2569 → ['เมย69']"""
    abbr = _MONTH_SHEET_MAP.get(month, "")
    yr2 = year_be - 2500
    return [f"{abbr}{yr2}", f"{abbr} {yr2}", f"{abbr}{year_be}", f"{abbr}.{yr2}"]


def read_cm_personnel(start_date: date, end_date: date) -> dict:
    """อ่าน CM personnel attendance ในช่วง start–end
    คืน: {
      "personnel": [{"no": 1, "type": "วิศวกรโยธา/วิศวกรโครงการ", "name": "นายโชคพิพัฒน์...", "attendance": [1, 1, "-", ...]}],
      "totals": [6, 6, 8, ...]   # รวมต่อวัน
    }
    """
    if not os.path.exists(DATA_CM):
        return {"personnel": [], "totals": []}

    wb = openpyxl.load_workbook(DATA_CM, data_only=True)

    # สร้าง list ของวันที่ใน range
    dates = []
    d = start_date
    while d <= end_date:
        dates.append(d)
        d += timedelta(days=1)

    # อ่านแต่ละวันจาก sheet ของเดือนนั้น
    # โครงสร้าง sheet:
    # Row 1-4: header
    # Row 5: ลำดับที่ | ประเภทบุคคล | ชื่อ-สกุล | (empty) | เดือน...
    # Row 6: (cont) | (cont) | (cont) | (empty) | 1 | 2 | 3 | ... | 30/31
    # Row 7-15: data (9 rows)
    # Row 16: รวม

    # collect personnel info จาก sheet แรกที่มี
    personnel = []
    for d_iter in dates:
        sheet_candidates = _sheet_name_for_month(d_iter.month, d_iter.year + 543)
        ws = None
        for cand in sheet_candidates:
            if cand in wb.sheetnames:
                ws = wb[cand]
                break
        if ws is not None and not personnel:
            for ri in range(7, 16):
                no = ws.cell(ri, 2).value
                ptype = ws.cell(ri, 3).value
                pname = ws.cell(ri, 4).value
                if no is None and ptype is None:
                    continue
                personnel.append({
                    "no": no, "type": str(ptype or "").strip(),
                    "name": str(pname or "").strip(),
                    "attendance": []
                })
            break

    if not personnel:
        return {"personnel": [], "totals": []}

    # อ่าน attendance ทุกวัน
    totals = []
    for d_iter in dates:
        sheet_candidates = _sheet_name_for_month(d_iter.month, d_iter.year + 543)
        ws = None
        for cand in sheet_candidates:
            if cand in wb.sheetnames:
                ws = wb[cand]
                break
        if ws is None:
            for p in personnel:
                p["attendance"].append("-")
            totals.append(0)
            continue
        # col index ของวันนี้: day 1 อยู่ที่ col 6 (F), day 16 = col 21, day 31 = col 36
        col = 5 + d_iter.day  # day 1 → col 6 (F)
        day_total = 0
        for pi, p in enumerate(personnel):
            row_idx = 7 + pi  # personnel row
            v = ws.cell(row_idx, col).value
            if isinstance(v, (int, float)) and v >= 1:
                p["attendance"].append("1")
                day_total += int(v)
            elif v in (None, "", "-"):
                p["attendance"].append("-")
            else:
                p["attendance"].append(str(v))
        totals.append(day_total)

    return {"personnel": personnel, "totals": totals}


def fill_appendix4_xlsx(template_path: str, week_no: int,
                        start_date: date, end_date: date, cm_data: dict) -> bytes:
    """เติมไฟล์ ภาคผนวก 4 xlsx (CM personnel weekly)"""
    wb = openpyxl.load_workbook(template_path)
    ws = wb.active  # การปฏิบัติงานของผู้ให้บริการ

    year_be = start_date.year + 543
    # update title row 5: "สัปดาห์ที่ XX/YYYY (วันที่ DD - DD เดือน YYYY)"
    months_thai = ["", "มกราคม", "กุมภาพันธ์", "มีนาคม", "เมษายน", "พฤษภาคม", "มิถุนายน",
                   "กรกฎาคม", "สิงหาคม", "กันยายน", "ตุลาคม", "พฤศจิกายน", "ธันวาคม"]
    month_name = months_thai[end_date.month]
    title = (f"สัปดาห์ที่ {week_no}/{year_be} "
             f"(วันที่ {start_date.day} – {end_date.day} {month_name} {year_be})")
    ws.cell(5, 5).value = title

    # update day numbers row 6: cols E..L (5..12)
    n_days = (end_date - start_date).days + 1
    for i in range(min(n_days, 8)):
        day = (start_date + timedelta(days=i)).day
        ws.cell(6, 5 + i).value = day

    # fill personnel data rows 7-15
    personnel = cm_data.get("personnel", [])
    for pi, p in enumerate(personnel[:9]):
        row_idx = 7 + pi
        for di in range(min(n_days, 8)):
            v = p["attendance"][di] if di < len(p["attendance"]) else "-"
            ws.cell(row_idx, 5 + di).value = v

    # totals row 16
    totals = cm_data.get("totals", [])
    for di in range(min(n_days, 8)):
        ws.cell(16, 5 + di).value = totals[di] if di < len(totals) else 0

    # ━━━━━━ บังคับให้ตารางพอดี 1 หน้ากระดาษ A4 แนวนอน ━━━━━━
    ws.page_setup.orientation = ws.ORIENTATION_LANDSCAPE
    ws.page_setup.paperSize = ws.PAPERSIZE_A4
    ws.page_setup.fitToWidth = 1
    ws.page_setup.fitToHeight = 1
    ws.sheet_properties.pageSetUpPr.fitToPage = True
    # ตั้ง print area ให้ครอบคลุมข้อมูลทั้งหมด (B1:L19)
    last_col = openpyxl.utils.get_column_letter(min(5 + max(n_days, 8) - 1, ws.max_column))
    ws.print_area = f"B1:{last_col}19"
    # margins แคบลงเพื่อให้พื้นที่พิมพ์เยอะขึ้น
    ws.page_margins.left = 0.3
    ws.page_margins.right = 0.3
    ws.page_margins.top = 0.4
    ws.page_margins.bottom = 0.4
    ws.page_margins.header = 0.2
    ws.page_margins.footer = 0.2
    # จัดกลางหน้ากระดาษ
    ws.print_options.horizontalCentered = True

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()
