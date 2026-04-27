"""
report_generator.py v3
รายงานประจำวัน: เติมข้อมูลลงใน template_daily.docx แทนสร้างใหม่
"""

import io, json, re, os, httpx
from copy import deepcopy
from datetime import datetime, date, timedelta
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

THAI_MONTHS_FULL = ["","มกราคม","กุมภาพันธ์","มีนาคม","เมษายน","พฤษภาคม","มิถุนายน",
                    "กรกฎาคม","สิงหาคม","กันยายน","ตุลาคม","พฤศจิกายน","ธันวาคม"]

_MONTHS_PAT = "|".join(THAI_MONTHS_FULL[1:])

def clean_caption(text: str) -> str:
    """ตัดบรรทัดวันที่ กำลังพล และเครื่องจักรออกจาก caption ใต้ภาพ"""
    if not text:
        return ""
    lines = []
    for line in text.split('\n'):
        s = line.strip()
        if not s:
            continue
        # ตัดบรรทัดวันที่ เช่น "วันที่ 22 เมษายน 2569 อากาศแจ่มใส"
        if re.search(rf'(?:วันที่).*(?:{_MONTHS_PAT}).*\d{{4}}', s):
            continue
        # ตัดบรรทัดกำลังพล เช่น "วิศวกร 2 คน หัวหน้าคนงาน 3 คน รวม 7 คน"
        if re.search(r'(?:วิศวกร|หัวหน้าคนงาน|หัวหน้า|กรรมกร|ช่างฝีมือ|ช่าง|คนงาน)\s*\d+\s*คน', s):
            continue
        # ตัดบรรทัดเครื่องจักร เช่น "รถแบ็คโฮ 1 คัน", "รถนํ้า 1 คัน"
        s_norm = s.replace('นํ้า', 'น้ำ')
        if re.search(r'(?:รถแบ็คโฮ|แบ็คโฮ|รถขุด|รถบรรทุก|รถเครน|รถบด|รถน้ำ|รถเกรด|รถสูบน้ำ|รถแทร็กเตอร์|กล้องสำรวจ|กล้องระดับ|เครื่องจี้|เครื่องเชื่อม|เครื่องตบ|เครื่องสูบ)\s*\d+\s*(?:คัน|ค้น|ค่น|ตัว|เครื่อง)', s_norm):
            continue
        # ตัดบรรทัดระดับน้ำ เช่น "+92.50", "ระดับน้ำ +92.60", "ระดับนํ้า +92.60"
        if re.match(r'^[+-]\d+(?:\.\d+)?\s*(?:ม\.|เมตร|m)?$', s):
            continue
        if re.search(r'ระดับน[^\s]*า', s) or re.search(r'ระดับน้ำ', s):
            continue
        lines.append(s)
    result = "\n".join(lines)
    # แก้ "วันที่ วันที่" ซ้ำ กรณีที่หลุดผ่านมา
    result = re.sub(r'วันที่\s+วันที่', 'วันที่', result)
    return result

_THAI_DIGITS_MAP = str.maketrans("0123456789", "๐๑๒๓๔๕๖๗๘๙")

def to_thai_digits(s) -> str:
    """แปลงเลขอารบิก → เลขไทย: 24 → ๒๔"""
    return str(s).translate(_THAI_DIGITS_MAP)

def thai_date(d):
    if isinstance(d, str): d = date.fromisoformat(d)
    return f"{d.day} {THAI_MONTHS_FULL[d.month]} {d.year+543}"

def thai_date_th_digits(d):
    """24 เมษายน 2569 → ๒๔ เมษายน ๒๕๖๙"""
    if isinstance(d, str): d = date.fromisoformat(d)
    return f"{to_thai_digits(d.day)} {THAI_MONTHS_FULL[d.month]} {to_thai_digits(d.year+543)}"

def thai_date_short(d):
    if isinstance(d, str): d = date.fromisoformat(d)
    return f"{d.day:02d}/{d.month:02d}/{str(d.year+543)[2:]}"

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

def add_header_row(table, headers, bg="1F4E79"):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = ""
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        run.font.size = Pt(11)

def style_doc(doc):
    s = doc.styles["Normal"]
    s.font.name = "TH Sarabun New"
    s.font.size = Pt(14)
    for sec in doc.sections:
        sec.page_height = Cm(29.7); sec.page_width = Cm(21.0)
        sec.left_margin = Cm(2.5);  sec.right_margin = Cm(2.0)
        sec.top_margin  = Cm(2.0);  sec.bottom_margin = Cm(2.0)

def add_title_block(doc, title, subtitle, project_name):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title); r.bold = True; r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x1F,0x4E,0x79)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(subtitle).font.size = Pt(13)
    if project_name:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r3 = p3.add_run(f"โครงการ: {project_name}")
        r3.bold = True; r3.font.size = Pt(13)
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("─"*60); r4.font.size = Pt(10)
    r4.font.color.rgb = RGBColor(0x1F,0x4E,0x79)
    doc.add_paragraph()

def add_image_caption(doc, caption: str):
    """เพิ่ม caption ใต้ภาพโดยใช้ line break แทน paragraph break เพื่อไม่ให้มีช่องว่าง
    บังคับใช้ TH SarabunIT๙ ทั้ง ascii + cs (complex script ภาษาไทย)
    """
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_before = Pt(4)
    cap_p.paragraph_format.space_after = Pt(6)
    lines = [l for l in caption.split('\n') if l.strip()]
    for i, line in enumerate(lines):
        run = cap_p.add_run(line)
        run.font.name = "TH SarabunIT๙"
        run.font.size = Pt(16)
        # set CS font (สำคัญ! python-docx default ไม่ตั้งให้)
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        for attr in ("ascii", "hAnsi", "cs"):
            rFonts.set(qn(f"w:{attr}"), "TH SarabunIT๙")
        # ตั้ง CS size ด้วย (Thai รัน complex script)
        szCs = rPr.find(qn("w:szCs"))
        if szCs is None:
            szCs = OxmlElement("w:szCs")
            rPr.append(szCs)
        szCs.set(qn("w:val"), "32")  # 32 half-points = 16pt
        if i < len(lines) - 1:
            run.add_break()


async def download_image_bytes(url):
    try:
        async with httpx.AsyncClient(timeout=15) as c:
            r = await c.get(url); r.raise_for_status(); return r.content
    except Exception as e:
        print(f"⚠️ img download: {e}"); return None


def add_labor_table(doc, data: dict):
    """ตารางกำลังพล: วิศวกร / หัวหน้า / ช่าง / กรรมกร / รวม"""
    engineers       = data.get("engineers") or 0
    foremen         = data.get("foremen") or 0
    skilled_workers = data.get("skilled_workers") or 0
    laborers        = data.get("laborers") or 0
    total           = data.get("total_workers") or (engineers+foremen+skilled_workers+laborers)

    if total == 0 and not any([engineers, foremen, skilled_workers, laborers]):
        p = doc.add_paragraph("— ไม่มีข้อมูลกำลังพล —")
        p.runs[0].font.size = Pt(12)
        return

    tbl = doc.add_table(rows=2, cols=5)
    tbl.style = "Table Grid"
    add_header_row(tbl, ["วิศวกร/ช่าง (คน)", "หัวหน้าคนงาน (คน)", "ช่างฝีมือ (คน)", "กรรมกร (คน)", "รวม (คน)"])
    row = tbl.rows[1]
    for i, v in enumerate([engineers, foremen, skilled_workers, laborers, total]):
        row.cells[i].text = str(v) if v else "—"
        row.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if v and i == 4:  # รวม — สีเน้น
            set_cell_bg(row.cells[i], "D6E4F0")
            if row.cells[i].paragraphs[0].runs:
                row.cells[i].paragraphs[0].runs[0].bold = True


def add_equipment_table(doc, data: dict):
    """ตารางเครื่องจักร"""
    equip_raw = data.get("equipment")
    if isinstance(equip_raw, str):
        try: equipment = json.loads(equip_raw)
        except: equipment = []
    else:
        equipment = equip_raw or []

    if not equipment:
        doc.add_paragraph("— ไม่มีข้อมูลเครื่องจักร —").runs[0].font.size = Pt(12)
        return

    tbl = doc.add_table(rows=len(equipment)+1, cols=3)
    tbl.style = "Table Grid"
    add_header_row(tbl, ["ลำดับ", "ประเภทเครื่องจักร/ยานพาหนะ", "จำนวน"])
    for i, eq in enumerate(equipment):
        row = tbl.rows[i+1]
        row.cells[0].text = str(i+1)
        row.cells[1].text = eq.get("name", "—")
        row.cells[2].text = f"{eq.get('qty',1)} {eq.get('unit','คัน')}"
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i % 2 == 0:
            for cell in row.cells: set_cell_bg(cell, "EBF3FB")


def add_signature_table(doc, left="ผู้รายงาน (Reported by)", right="ผู้ตรวจสอบ (Checked by)"):
    tbl = doc.add_table(rows=3, cols=2)
    tbl.style = "Table Grid"
    tbl.rows[0].cells[0].text = left
    tbl.rows[0].cells[1].text = right
    tbl.rows[1].cells[0].text = "\n\n"
    tbl.rows[1].cells[1].text = "\n\n"
    tbl.rows[2].cells[0].text = "วันที่: ____________________"
    tbl.rows[2].cells[1].text = "วันที่: ____________________"
    for cell in tbl.rows[0].cells:
        set_cell_bg(cell, "D6E4F0")
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True


# ════════════════════════════════════════
# TEMPLATE HELPERS (DAILY REPORT)
# ════════════════════════════════════════

TEMPLATE_DAILY = os.path.join(os.path.dirname(os.path.abspath(__file__)), "template_daily.docx")

# โครงสร้าง paragraph ใน template_daily.docx:
#  [1]  วันที่ {day} เดือน {month} พ.ศ.{year}   → run[2]=day, run[6]=month+year
#  [11] ค่าระดับน้ำขึ้นสูงสุด {+xx.xx ม.}        → run[2]=water level
#  [12] งานที่ทำ  1. {activity1}                → run[3]=text
#  [13] \t\t2. {activity2}                      → run[2]=text
#  [20] วิศวกร/ช่าง {x} ... หัวหน้าคนงาน {x} กรรมกร {x} รวม {x}
#       → run[2]=eng, run[7]=foremen, run[11]=laborers(padded), run[15]=total
#  [22-27] ตารางเครื่องจักร  → regex replace qty

_EQUIP_MAP = {
    22: ["รถเฮี้ยบ", "รถเทเลอร์", "รถสกัดคอนกรีต", "รถเกรด"],
    23: ["รถแบ็คโฮ", "รถบด", "รถน้ำ", "รถบรรทุก", "รถแทร็กเตอร์"],
    24: ["รถเครน", "รถบดล้อยาง", "ปั่นจั่น", "รถสกัดคอนกรีตเสาเข็ม"],
    25: ["รถบรรทุก 10 ล้อ", "รถบรรทุก 6 ล้อ", "กล้องสำรวจแนว", "กล้องระดับ"],
    26: ["เครื่องจี้คอนกรีต", "เครื่องเชื่อม", "เครื่องตบดิน", "เครื่องสูบน้ำ"],
    27: ["รถบดสั่นสะเทือน", "รถน้ำ", "รถPRIME COAT", "รถPAVE ยาง"],
}


def _tpl_set_run(para, idx, text):
    if idx < len(para.runs):
        para.runs[idx].text = text


# จำนวน trailing tab ที่พอดีกับ content width โดยไม่ wrap:
# content = 9026 twips, default tab = 720 twips → max 12 stops จาก 0
# para มี leading tabs 2 ตัว (1440 twips) → trailing ≤ 10 ก็ปลอดภัย
_TARGET_TRAILING_TABS = 10

# Wingdings checkbox characters
_CKBOX_CHECKED   = ("F0FE", "Wingdings")    # ☑
_CKBOX_UNCHECKED = ("F0A3", "Wingdings 2")  # ☐

# weather_morning value → (para_idx, run_idx ที่มี w:sym)
_WEATHER_MAP = {
    "แจ่มใส":        (8, 1),
    "เมฆมาก":       (8, 8),
    "ฝนตกเล็กน้อย": (9, 1),
    "ฝนตกหนัก":     (9, 5),
}

def _tpl_set_checkbox(run, checked: bool):
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    sym = run._r.find(f"{{{ns}}}sym")
    if sym is None:
        return
    char, font = _CKBOX_CHECKED if checked else _CKBOX_UNCHECKED
    sym.set(qn("w:char"), char)
    sym.set(qn("w:font"), font)


def _tpl_fill_weather(paras, weather_morning: str):
    """ติ๊กช่องสภาพอากาศตาม weather_morning — untick ทั้งหมดก่อน แล้วติ๊กที่ตรง"""
    if not weather_morning:
        return
    # untick ทั้งหมด
    for pidx, ridx in _WEATHER_MAP.values():
        _tpl_set_checkbox(paras[pidx].runs[ridx], False)
    # หา key ที่ตรง
    matched = None
    if "ฝนตกหนัก" in weather_morning:
        matched = "ฝนตกหนัก"
    elif any(k in weather_morning for k in ("ฝนตกเล็กน้อย", "ฝนตกน้อย", "ฝนเล็กน้อย")):
        matched = "ฝนตกเล็กน้อย"
    elif any(k in weather_morning for k in ("ฝนตก", "ฝน")):
        matched = "ฝนตกเล็กน้อย"  # ฝนทั่วไป = เล็กน้อย
    elif any(k in weather_morning for k in ("เมฆมาก", "มืดครึ้ม", "ครึ้ม", "มีเมฆ")):
        matched = "เมฆมาก"
    elif any(k in weather_morning for k in ("แจ่มใส", "แดด", "ร้อน", "หมอก")):
        matched = "แจ่มใส"
    if matched:
        pidx, ridx = _WEATHER_MAP[matched]
        _tpl_set_checkbox(paras[pidx].runs[ridx], True)


def _remove_underline_from_para(para):
    """ลบ <w:u> ออกจากทุก run ใน paragraph — ใช้กับ activity rows เพื่อเอาเส้นประออก"""
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    for run in para.runs:
        rPr = run._r.find(f"{{{ns}}}rPr")
        if rPr is not None:
            for u_el in rPr.findall(f"{{{ns}}}u"):
                rPr.remove(u_el)


def _set_tight_spacing(p_elem):
    """ตั้ง spacing ให้ชิดติดกัน: before=0, after=0, single line — ใช้กับ activity rows"""
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    pPr = p_elem.find(f"{{{ns}}}pPr")
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_elem.insert(0, pPr)
    # ลบ spacing เก่าแล้วใส่ใหม่
    for sp in pPr.findall(f"{{{ns}}}spacing"):
        pPr.remove(sp)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "240")
    spacing.set(qn("w:lineRule"), "auto")
    # ใส่ spacing ก่อน jc (หรือต่อท้ายถ้าไม่มี jc)
    jc = pPr.find(f"{{{ns}}}jc")
    if jc is not None:
        pPr.insert(list(pPr).index(jc), spacing)
    else:
        pPr.append(spacing)


def _tpl_set_activity_line(para, text_run_idx, text):
    """ใส่ข้อความ activity, ลบ trailing tabs ทั้งหมดออก (XML level), และลบเส้นประ"""
    if text_run_idx < len(para.runs):
        para.runs[text_run_idx].text = text
    # ลบ <w:r> elements ทั้งหมดหลัง text_run_idx ออกจาก XML โดยตรง
    # (run.text="" ไม่พอ เพราะ <w:tab/> ยังคงอยู่ใน XML)
    runs_to_remove = [para.runs[i]._r for i in range(text_run_idx + 1, len(para.runs))]
    for r_el in runs_to_remove:
        r_el.getparent().remove(r_el)
    # ลบ custom tab stops
    pPr = para._p.get_or_add_pPr()
    for tabs_el in pPr.findall(qn("w:tabs")):
        pPr.remove(tabs_el)
    # ลบเส้นประ (dotted underline) ออกจากทุก run ที่เหลือ
    _remove_underline_from_para(para)
    # ตั้ง spacing ให้ชิดติดกัน (ไม่มี space before/after, single line)
    _set_tight_spacing(para._p)


def _tpl_rebuild_para(para, new_text):
    """ล้าง runs ทั้งหมด ใส่ new_text ใน run แรก"""
    for i, run in enumerate(para.runs):
        run.text = new_text if i == 0 else ""


def _tpl_delete_para(para):
    p = para._element
    p.getparent().remove(p)


def _tpl_insert_activity(ref_elem, num, act_text):
    """แทรก paragraph ของ activity ถัดจาก ref_elem, คืนค่า element ใหม่
    deepcopy para[13] structure แล้วลบ trailing tabs, dotted underline และ set spacing
    """
    ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    new_p = deepcopy(ref_elem)
    ref_elem.addnext(new_p)

    # ตั้งข้อความและลบ trailing tab runs ออกจาก XML
    t_list = new_p.findall(f".//{{{ns}}}t")
    if t_list:
        t_list[0].text = f"{num}. {act_text}"
        t_list[0].set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        # หา <w:r> ที่เป็น parent ของ <w:t> แล้วลบ <w:r> ทั้งหมดหลังจากนั้น
        text_run_elem = t_list[0].getparent()
        all_runs = new_p.findall(f"{{{ns}}}r")
        found = False
        for r_el in all_runs:
            if found:
                r_el.getparent().remove(r_el)
            elif r_el is text_run_elem:
                found = True

    # ลบ dotted underline ออกจากทุก run ที่เหลือ
    for rPr in new_p.findall(f".//{{{ns}}}rPr"):
        for u_el in rPr.findall(f"{{{ns}}}u"):
            rPr.remove(u_el)
    # ตั้ง spacing ให้ชิดติดกัน
    _set_tight_spacing(new_p)
    return new_p


def _tpl_fill_equipment(paras, daily_data):
    equip_raw = daily_data.get("equipment")
    if isinstance(equip_raw, str):
        try: equipment = json.loads(equip_raw)
        except: equipment = []
    else:
        equipment = equip_raw or []

    lookup = {eq.get("name", "").strip(): eq.get("qty", 1) for eq in equipment}

    for pidx, names in _EQUIP_MAP.items():
        para = paras[pidx]
        text = para.text
        for name in names:
            qty = lookup.get(name)
            pattern = re.escape(name) + r"[…\.\d]+(?:คัน|ตัว|เครื่อง)"

            def _repl(m, n=name, q=qty):
                unit = re.search(r"(?:คัน|ตัว|เครื่อง)$", m.group()).group()
                return f"{n}…{q}….{unit}" if q else f"{n}………{unit}"

            text = re.sub(pattern, _repl, text)
        _tpl_rebuild_para(para, text)


# ════════════════════════════════════════
# DAILY REPORT
# ════════════════════════════════════════

async def generate_daily(work_date: str, daily_data: dict, project_name: str = "โครงการก่อสร้าง",
                         include_images: bool = True) -> bytes:
    doc = Document(TEMPLATE_DAILY)
    paras = doc.paragraphs
    d = date.fromisoformat(work_date)

    # 1. วันที่
    _tpl_set_run(paras[1], 2, str(d.day))
    _tpl_set_run(paras[1], 6, f"{THAI_MONTHS_FULL[d.month]}   พ.ศ.{d.year+543}")

    # 2. สภาพอากาศ — ติ๊ก checkbox ให้ตรงกับ weather_morning
    _tpl_fill_weather(paras, daily_data.get("weather_morning") or "")

    # 3. ระดับน้ำ
    wl = daily_data.get("water_level")
    wl_str = ((f"+{wl:.2f}" if wl >= 0 else f"{wl:.2f}") + " ม.") if wl is not None else "— ม."
    _tpl_set_run(paras[11], 2, wl_str)

    # 4. งานที่ทำ
    activities = daily_data.get("activities") or []
    act1 = activities[0].get("desc") or activities[0].get("description") if activities else ""
    _tpl_set_activity_line(paras[12], text_run_idx=3,
                           text=f"1. {act1}" if act1 else "1. —")

    if len(activities) >= 2:
        act2 = activities[1].get("desc") or activities[1].get("description") or ""
        _tpl_set_activity_line(paras[13], text_run_idx=2, text=f"2. {act2}")
        ref_elem = paras[13]._element
        for i, act in enumerate(activities[2:], start=3):
            act_text = act.get("desc") or act.get("description") or ""
            ref_elem = _tpl_insert_activity(ref_elem, i, act_text)
    else:
        # มีงานแค่ 1 รายการ → ลบ para[13] ทิ้งเพื่อให้ ปัญหาอุปสรรค ต่อเนื่องทันที
        _tpl_delete_para(paras[13])

    # 4. กำลังพล
    eng = (daily_data.get("engineers") or 0) + (daily_data.get("skilled_workers") or 0)
    fmn = daily_data.get("foremen") or 0
    lab = daily_data.get("laborers") or 0
    total = daily_data.get("total_workers") or (eng + fmn + lab)

    pw = paras[20]
    _tpl_set_run(pw, 2, str(eng) if eng else "—")
    _tpl_set_run(pw, 7, str(fmn) if fmn else "—")
    if 11 < len(pw.runs):
        pw.runs[11].text = f"      {lab}     " if lab else "      —     "
    _tpl_set_run(pw, 15, str(total) if total else "—")

    # 5. เครื่องจักร
    _tpl_fill_equipment(paras, daily_data)

    # 6. รูปภาพ — ลบ placeholder ใน template แล้วใส่รูปจริง
    # (include_images=False → ใช้ตอน embed ใน weekly report เพราะมีภาคผนวก 1 ภาพถ่ายแยกอยู่แล้ว)
    heading_para = next((p for p in doc.paragraphs if "รูปภาพประกอบ" in p.text), None)
    if heading_para:
        # บังคับ font ของ "รูปภาพประกอบ" เป็น TH SarabunIT๙
        for run in heading_para.runs:
            run.font.name = "TH SarabunIT๙"
            rPr = run._r.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                from docx.oxml import OxmlElement
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)
            for attr in ("ascii", "hAnsi", "cs"):
                rFonts.set(qn(f"w:{attr}"), "TH SarabunIT๙")
        heading_elem = heading_para._element
        found = False
        to_delete = []
        for p in doc.paragraphs:
            if found:
                to_delete.append(p)
            if p._element is heading_elem:
                found = True
        for p in to_delete:
            _tpl_delete_para(p)
        # ถ้าไม่ต้องการรูป → ลบหัวข้อ "รูปภาพประกอบ" ทิ้งด้วย
        if not include_images:
            _tpl_delete_para(heading_para)

    images = daily_data.get("images") or []
    if include_images and images:
        for img_info in images:
            url = img_info.get("url") or img_info.get("image_url")
            acts_text = clean_caption(img_info.get("caption") or "")
            # ใช้เลขไทยสำหรับวันที่
            caption = f"วันที่ {thai_date_th_digits(d)}"
            if acts_text:
                caption += f"\n{acts_text}"
            if not url:
                continue
            img_bytes = await download_image_bytes(url)
            if img_bytes:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(io.BytesIO(img_bytes), width=Inches(5.5))
            add_image_caption(doc, caption)
            doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ════════════════════════════════════════
# WEEKLY REPORT
# ════════════════════════════════════════

async def generate_weekly(week_start: str, daily_list: list, project_name: str = "โครงการก่อสร้าง",
                          week_no: int = None, week_end: str = None) -> bytes:
    doc = Document()
    style_doc(doc)
    ws = date.fromisoformat(week_start)
    we = date.fromisoformat(week_end) if week_end else ws + timedelta(days=6)
    if week_no:
        subtitle = (f"สัปดาห์ที่ {week_no}  |  {thai_date(ws)} — {thai_date(we)}"
                    f"  |  ฉบับที่ WR-{ws.strftime('%Y%m')}-W{week_no}")
    else:
        subtitle = f"{thai_date(ws)} — {thai_date(we)}  |  ฉบับที่ WR-{ws.strftime('%Y%m%d')}"
    add_title_block(doc, "รายงานความก้าวหน้าประจำสัปดาห์ (WEEKLY PROGRESS REPORT)",
        subtitle, project_name)

    # สรุปสัปดาห์
    doc.add_heading("1. สรุปภาพรวมสัปดาห์", level=2)
    total_w = sum(d.get("total_workers") or 0 for d in daily_list)
    total_e = sum(d.get("engineers") or 0 for d in daily_list)
    total_f = sum(d.get("foremen") or 0 for d in daily_list)
    total_s = sum(d.get("skilled_workers") or 0 for d in daily_list)
    total_l = sum(d.get("laborers") or 0 for d in daily_list)

    sum_tbl = doc.add_table(rows=2, cols=5)
    sum_tbl.style = "Table Grid"
    add_header_row(sum_tbl, ["วันทำงาน","วิศวกร (คน-วัน)","หัวหน้า (คน-วัน)","ช่าง (คน-วัน)","กรรมกร (คน-วัน)"])
    r = sum_tbl.rows[1]
    for i,v in enumerate([f"{len(daily_list)} วัน", str(total_e), str(total_f), str(total_s), str(total_l)]):
        r.cells[i].text = v
        r.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ตารางสรุปรายวัน
    doc.add_heading("2. ตารางสรุปงานรายวัน", level=2)
    day_tbl = doc.add_table(rows=len(daily_list)+1, cols=5)
    day_tbl.style = "Table Grid"
    add_header_row(day_tbl, ["วันที่","อากาศ","คนงานรวม (คน)","กิจกรรมหลัก","เครื่องจักร"])
    for i, d in enumerate(daily_list):
        acts = d.get("activities") or []
        act_str = ", ".join(a.get("desc") or a.get("description","") for a in acts[:2])
        if len(acts)>2: act_str += f" (+{len(acts)-2})"

        equip_raw = d.get("equipment")
        if isinstance(equip_raw, str):
            try: equip = json.loads(equip_raw)
            except: equip = []
        else: equip = equip_raw or []
        eq_str = ", ".join(f"{e['name']} {e['qty']}{e['unit']}" for e in equip[:2]) or "—"

        row = day_tbl.rows[i+1]
        row.cells[0].text = thai_date_short(d.get("work_date",""))
        row.cells[1].text = d.get("weather_morning") or "—"
        row.cells[2].text = str(d.get("total_workers") or "—")
        row.cells[3].text = act_str or "—"
        row.cells[4].text = eq_str
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i%2==0:
            for cell in row.cells: set_cell_bg(cell,"EBF3FB")
    doc.add_paragraph()

    # รูปภาพ + caption รายวัน
    doc.add_heading("3. ภาพความก้าวหน้าประจำสัปดาห์", level=2)
    for day_data in daily_list:
        images = day_data.get("images") or []
        acts   = day_data.get("activities") or []
        if not images and not acts: continue

        wp = doc.add_paragraph()
        wr = wp.add_run(f"▶  {thai_date(day_data.get('work_date',''))}")
        wr.bold = True; wr.font.size = Pt(14)
        wr.font.color.rgb = RGBColor(0x1F,0x4E,0x79)

        # กำลังพลและเครื่องจักรของวัน
        labor_parts = []
        if day_data.get("engineers"):       labor_parts.append(f"วิศวกร {day_data['engineers']} คน")
        if day_data.get("foremen"):         labor_parts.append(f"หัวหน้า {day_data['foremen']} คน")
        if day_data.get("skilled_workers"): labor_parts.append(f"ช่าง {day_data['skilled_workers']} คน")
        if day_data.get("laborers"):        labor_parts.append(f"กรรมกร {day_data['laborers']} คน")
        if labor_parts:
            lp = doc.add_paragraph(f"👷 {', '.join(labor_parts)}")
            lp.runs[0].font.size = Pt(12)

        if images:
            for img_info in images:
                url     = img_info.get("url") or img_info.get("image_url")
                acts_text = clean_caption(img_info.get("caption") or "")
                _wd = day_data.get("work_date", "")
                caption = f"วันที่ {thai_date(_wd)}" if _wd else ""
                if acts_text:
                    caption += f"\n{acts_text}" if caption else acts_text
                if not url: continue
                img_bytes = await download_image_bytes(url)
                if img_bytes:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(io.BytesIO(img_bytes), width=Inches(5.0))
                add_image_caption(doc, caption)
        elif acts:
            for act in acts:
                bp = doc.add_paragraph(act.get("desc") or act.get("description",""), style="List Bullet")
        doc.add_paragraph()

    doc.add_heading("ลงชื่อ / Signature", level=2)
    add_signature_table(doc, "ผู้รับจ้าง (Contractor)", "ผู้ควบคุมงาน (Inspector)")

    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()


# ════════════════════════════════════════
# MONTHLY REPORT
# ════════════════════════════════════════

async def generate_monthly(month_str: str, daily_list: list, project_name: str = "โครงการก่อสร้าง") -> bytes:
    doc = Document()
    style_doc(doc)
    yr, mo = int(month_str[:4]), int(month_str[5:7])
    month_label = f"{THAI_MONTHS_FULL[mo]} {yr+543}"
    add_title_block(doc, "รายงานความก้าวหน้าประจำเดือน (MONTHLY PROGRESS REPORT)",
        f"{month_label}  |  ฉบับที่ MPR-{month_str.replace('-','')}", project_name)

    # สรุปเดือน
    doc.add_heading("1. สรุปภาพรวมประจำเดือน", level=2)
    total_w = sum(d.get("total_workers") or 0 for d in daily_list)
    sum_tbl = doc.add_table(rows=2, cols=4)
    sum_tbl.style = "Table Grid"
    add_header_row(sum_tbl,["วันทำงาน","คนงานรวม (คน-วัน)","กิจกรรมรวม","รูปภาพรวม"])
    r = sum_tbl.rows[1]
    for i,v in enumerate([f"{len(daily_list)} วัน", str(total_w),
                           str(sum(len(d.get("activities") or []) for d in daily_list)),
                           str(sum(len(d.get("images") or []) for d in daily_list))]):
        r.cells[i].text = v
        r.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ตารางสรุปรายวัน
    doc.add_heading("2. ตารางสรุปงานประจำเดือน", level=2)
    day_tbl = doc.add_table(rows=len(daily_list)+1, cols=5)
    day_tbl.style = "Table Grid"
    add_header_row(day_tbl,["วันที่","อากาศ","คนงาน (คน)","กิจกรรมหลัก","เครื่องจักร"])
    for i,d in enumerate(daily_list):
        acts = d.get("activities") or []
        act_str = ", ".join(a.get("desc") or a.get("description","") for a in acts[:2])
        if len(acts)>2: act_str += f"(+{len(acts)-2})"

        equip_raw = d.get("equipment")
        if isinstance(equip_raw, str):
            try: equip = json.loads(equip_raw)
            except: equip = []
        else: equip = equip_raw or []
        eq_str = ", ".join(f"{e['name']} {e['qty']}{e['unit']}" for e in equip[:2]) or "—"

        row = day_tbl.rows[i+1]
        row.cells[0].text = thai_date_short(d.get("work_date",""))
        row.cells[1].text = d.get("weather_morning") or "—"
        row.cells[2].text = str(d.get("total_workers") or "—")
        row.cells[3].text = act_str or "—"
        row.cells[4].text = eq_str
        row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i%2==0:
            for cell in row.cells: set_cell_bg(cell,"EBF3FB")
    doc.add_paragraph()

    # รูปภาพ
    doc.add_heading("3. ภาพความก้าวหน้าประจำเดือน", level=2)
    for day_data in daily_list:
        images = day_data.get("images") or []
        if not images: continue
        dp = doc.add_paragraph()
        dr = dp.add_run(f"▶  {thai_date(day_data.get('work_date',''))}")
        dr.bold = True; dr.font.size = Pt(13)
        dr.font.color.rgb = RGBColor(0x1F,0x4E,0x79)
        for img_info in images:
            url     = img_info.get("url") or img_info.get("image_url")
            acts_text = clean_caption(img_info.get("caption") or "")
            _wd = day_data.get("work_date", "")
            caption = f"วันที่ {thai_date(_wd)}" if _wd else ""
            if acts_text:
                caption += f"\n{acts_text}" if caption else acts_text
            if not url: continue
            img_bytes = await download_image_bytes(url)
            if img_bytes:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(io.BytesIO(img_bytes), width=Inches(4.5))
            add_image_caption(doc, caption)
        doc.add_paragraph()

    doc.add_heading("ลงชื่อ / Signature", level=2)
    add_signature_table(doc, "ผู้รับจ้าง (Contractor)", "ผู้ควบคุมงาน (Inspector)")

    buf = io.BytesIO(); doc.save(buf); return buf.getvalue()
