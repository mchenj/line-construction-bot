"""
weekly_memo.py
เติมข้อมูลใน 2 ไฟล์:
- 00_memo.docx              (บันทึกข้อความ)
- 00_cover_letter.docx      (ใบปะหน้ารายงานประจำสัปดาห์)

Behavior:
- เลขที่หนังสือ + ลงวันที่ → เปลี่ยนเป็น 🔴 สีแดง (placeholder ให้ user แก้เอง)
- เลขสัปดาห์ + ช่วงวันที่ → ดึงจาก parameter
- ตารางบันทึกข้อความ (memo): ดึง %แผนงาน/%ผลงาน/%เร็วช้า/วันสะสม จาก Excel
"""

import os, io, re
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


_TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "templates_weekly")
TEMPLATE_MEMO         = os.path.join(_TEMPLATE_DIR, "00_memo.docx")
TEMPLATE_COVER_LETTER = os.path.join(_TEMPLATE_DIR, "00_cover_letter.docx")

THAI_MONTHS = ["", "มกราคม", "กุมภาพันธ์", "มีนาคม", "เมษายน", "พฤษภาคม", "มิถุนายน",
               "กรกฎาคม", "สิงหาคม", "กันยายน", "ตุลาคม", "พฤศจิกายน", "ธันวาคม"]


def _thai_date_full(d: date) -> str:
    """24 เมษายน 2569"""
    return f"{d.day} {THAI_MONTHS[d.month]} {d.year + 543}"


def _date_range_text(d1: date, d2: date) -> str:
    """16 – 23 เมษายน 2569 หรือ 28 มีนาคม – 4 เมษายน 2569 (กรณีข้ามเดือน)"""
    yr_th = d2.year + 543
    if d1.month == d2.month and d1.year == d2.year:
        return f"{d1.day} – {d2.day} {THAI_MONTHS[d2.month]} {yr_th}"
    return f"{d1.day} {THAI_MONTHS[d1.month]} – {d2.day} {THAI_MONTHS[d2.month]} {yr_th}"


# ════════════════════════════════════════
# Run-level helpers
# ════════════════════════════════════════

def _set_run_red(run):
    """ตั้งสีแดงให้ run"""
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)


def _replace_para_text(para, new_text: str, red: bool = False):
    """แทนข้อความใน paragraph โดยรักษา format ของ run แรก"""
    if not para.runs:
        run = para.add_run(new_text)
    else:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r._element.getparent().remove(r._element)
        run = para.runs[0]
    if red:
        _set_run_red(run)


def _redden_paragraph(para):
    """ทำทุก run ใน paragraph เป็นสีแดง"""
    for r in para.runs:
        _set_run_red(r)


def _replace_runs_by_pattern(para, pattern: str, replacement: str) -> bool:
    """หา substring ใน paragraph ที่ match regex pattern → แทนด้วย replacement
    รักษา runs อื่น (tabs, suffix text เช่น "จำนวน 12 ชุด") ไว้
    คืน True ถ้าแทนสำเร็จ
    """
    runs = list(para.runs)
    if not runs:
        return False
    # build full text + position map
    positions = []  # list of (start, end, run_idx)
    cum = 0
    parts = []
    for i, r in enumerate(runs):
        rt = r.text
        positions.append((cum, cum + len(rt), i))
        parts.append(rt)
        cum += len(rt)
    full_text = "".join(parts)
    m = re.search(pattern, full_text)
    if not m:
        return False
    s, e = m.start(), m.end()
    # หา start_run + end_run
    start_run = end_run = None
    s_off = e_off = 0
    for start, end, idx in positions:
        if start_run is None and start <= s < end:
            start_run = idx; s_off = s - start
        if start < e <= end:
            end_run = idx; e_off = e - start
            break
    if start_run is None or end_run is None:
        return False
    start_text = runs[start_run].text
    end_text = runs[end_run].text
    if start_run == end_run:
        # match ทั้งหมดอยู่ใน run เดียว
        runs[start_run].text = start_text[:s_off] + replacement + end_text[e_off:]
    else:
        runs[start_run].text = start_text[:s_off] + replacement
        runs[end_run].text = end_text[e_off:]
        # ลบ runs ระหว่าง start และ end
        for r in runs[start_run + 1 : end_run]:
            r._element.getparent().remove(r._element)
    return True


# ════════════════════════════════════════
# 00_memo.docx
# ════════════════════════════════════════

def fill_memo(week_no: int, week_start: date, week_end: date,
              progress: dict = None) -> bytes:
    """เติมข้อมูลใน 00_memo.docx
    Args:
        week_no: เลขสัปดาห์ (เช่น 89)
        week_start, week_end: ช่วงวันที่
        progress: dict จาก lookup_week_progress() — ถ้า None เว้นว่างไว้
    """
    doc = Document(TEMPLATE_MEMO)
    yr_th = week_start.year + 543
    date_range = _date_range_text(week_start, week_end)

    # P[2]: "ที่ KEC/หัวรอ/032/2569 (TAB) วันที่ 16 เมษายน 2569"
    # → เปลี่ยนเป็นสีแดงทั้งย่อหน้า (placeholder)
    # P[9]: "รายงานประจำสัปดาห์ ครั้งที่ 88 สัปดาห์ที่ 88/2569 (วันที่ 8 – 15 เม.ย. 2569)"
    for p in doc.paragraphs:
        text = p.text
        # หัวเลขที่หนังสือ + ลงวันที่ → ทำสีแดง
        if "ที่ KEC" in text or re.search(r"ที่\s+\S+/\S+/\d+", text):
            if "เรียน" not in text and "เรื่อง" not in text:
                _redden_paragraph(p)
                continue
        # บรรทัดรายงานประจำสัปดาห์ ครั้งที่ XX สัปดาห์ที่ XX/YYYY (วันที่ ...)
        if "รายงานประจำสัปดาห์" in text and "ครั้งที่" in text and "สัปดาห์ที่" in text:
            _replace_runs_by_pattern(
                p,
                r"รายงานประจำสัปดาห์\s*ครั้งที่\s*\d+\s*สัปดาห์ที่\s*\d+/25\d{2}\s*\(วันที่[^)]*\)",
                (f"รายงานประจำสัปดาห์ ครั้งที่ {week_no} "
                 f"สัปดาห์ที่ {week_no}/{yr_th} (วันที่ {date_range})")
            )

    # Table[0] — เติม progress data ในทุก row R3-R7 (5 rows ปีงบ)
    if progress and doc.tables:
        tbl = doc.tables[0]
        plan_cum   = f"{progress.get('plan_cum', 0):.2f}"
        actual_cum = f"{progress.get('actual_cum', 0):.2f}"
        diff_cum   = f"{progress.get('diff_cum', 0):.2f}"
        elapsed_d  = progress.get('elapsed_days', 0)
        elapsed_p  = progress.get('elapsed_pct', 0)
        elapsed_cell = f"{elapsed_d} วัน\n( {elapsed_p:.2f}%)"

        # rows 3-7 (data) — ปกติมี 5 ปีงบ
        for ri in range(3, min(len(tbl.rows), 10)):
            row = tbl.rows[ri]
            if len(row.cells) >= 4:
                _set_cell_text_keep_format(row.cells[0], plan_cum)
                _set_cell_text_keep_format(row.cells[1], actual_cum)
                _set_cell_text_keep_format(row.cells[2], diff_cum)
                _set_cell_text_keep_format(row.cells[3], elapsed_cell)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _set_cell_text_keep_format(cell, text: str):
    """แทน text ใน cell โดยรักษา font/style ของ run แรก รองรับ multiline"""
    if not cell.paragraphs:
        cell.add_paragraph(text)
        return
    p = cell.paragraphs[0]
    # ลบ paragraphs อื่น
    for old_p in cell.paragraphs[1:]:
        old_p._element.getparent().remove(old_p._element)
    lines = text.split("\n")
    if not p.runs:
        p.add_run(lines[0])
    else:
        # ใช้ first run
        p.runs[0].text = lines[0]
        # ลบ runs อื่นออก
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
    # บรรทัดถัดไป — เพิ่ม run ใหม่พร้อม break
    if len(lines) > 1:
        first_run = p.runs[0]
        for line in lines[1:]:
            new_run = p.add_run()
            # copy format
            try:
                new_run.font.name = first_run.font.name
                if first_run.font.size:
                    new_run.font.size = first_run.font.size
            except Exception:
                pass
            new_run.add_break()
            new_run.text = line


# ════════════════════════════════════════
# 00_cover_letter.docx
# ════════════════════════════════════════

def fill_cover_letter(week_no: int, week_start: date, week_end: date) -> bytes:
    """เติมข้อมูลใน 00_cover_letter.docx
    - P[0]: "ที่ CM67-09/115/2569" → สีแดง
    - P[1]: "16 เมษายน 2569" → สีแดง
    - บรรทัด "รายงานประจำสัปดาห์ที่ 88/2569 (วันที่ ...)" → แทนด้วย week_no + วันที่
    """
    doc = Document(TEMPLATE_COVER_LETTER)
    yr_th = week_start.year + 543
    date_range = _date_range_text(week_start, week_end)

    for i, p in enumerate(doc.paragraphs):
        text = p.text
        # P[0]: เลขที่หนังสือ → สีแดง
        if i == 0 and "ที่ " in text and "/" in text:
            _redden_paragraph(p)
            continue
        # P[1]: ลงวันที่ → สีแดง (มักเป็น "DD เดือน YYYY" สั้นๆ)
        if i == 1 and re.search(r"\d{1,2}\s+\S+\s+25\d{2}", text):
            _redden_paragraph(p)
            continue
        # บรรทัดที่มี "รายงานประจำสัปดาห์ที่ XX/YYYY (วันที่ ...)"
        # ใช้ run-level replace เพื่อรักษา tab + suffix "จำนวน 12 ชุด"
        if "รายงานประจำสัปดาห์ที่" in text and re.search(r"\d+/25\d{2}", text):
            _replace_runs_by_pattern(
                p,
                r"รายงานประจำสัปดาห์ที่\s*\d+/25\d{2}\s*\(วันที่[^)]*\)",
                f"รายงานประจำสัปดาห์ที่ {week_no}/{yr_th} (วันที่ {date_range})"
            )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
